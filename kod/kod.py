import win32evtlog
from datetime import datetime
import matplotlib.pyplot as plt
from collections import Counter
from docx import Document
import tkinter as tk
from tkinter import messagebox
from tkcalendar import DateEntry

# Olayları Çekme Fonksiyonu
def get_events(start_date, end_date):
    log_type = "System"
    server = "localhost"
    log_handle = win32evtlog.OpenEventLog(server, log_type)
    events = []
    flags = win32evtlog.EVENTLOG_FORWARDS_READ | win32evtlog.EVENTLOG_SEQUENTIAL_READ

    while True:
        records = win32evtlog.ReadEventLog(log_handle, flags, 0)
        if not records:
            break
        for event in records:
            event_time = event.TimeGenerated.Format()
            event_time = datetime.strptime(event_time, "%a %b %d %H:%M:%S %Y")
            if start_date <= event_time <= end_date:
                events.append({
                    "EventID": event.EventID,
                    "SourceName": event.SourceName,
                    "EventType": event.EventType,
                    "Message": event.StringInserts
                })

    win32evtlog.CloseEventLog(log_handle)
    return events

# Olayları Grafiğe Döken Fonksiyon
def plot_event_levels(events):
    levels = {1: "Error", 2: "Warning", 4: "Information"}
    event_levels = [levels.get(event["EventType"], "Other") for event in events]
    level_counts = Counter(event_levels)

    plt.figure(figsize=(8, 6))
    plt.bar(level_counts.keys(), level_counts.values(), color=['red', 'orange', 'blue'])
    plt.xlabel("Event Level")
    plt.ylabel("Count")
    plt.title("Event Levels Count")
    plt.show()

# Word'e Kaydetme Fonksiyonu
def save_to_word(events, filename="Event_Logs.docx"):
    doc = Document()
    doc.add_heading("Event Logs", 0)

    for event in events:
        doc.add_paragraph(f"Event ID: {event['EventID']}")
        doc.add_paragraph(f"Source Name: {event['SourceName']}")
        doc.add_paragraph(f"Level: {event['EventType']}")
        doc.add_paragraph(f"Message: {' '.join(event['Message']) if event['Message'] else 'N/A'}")
        doc.add_paragraph("\n-------------------\n")

    doc.save(filename)
    messagebox.showinfo("Kaydedildi", f"Word document saved as {filename}")

# İkinci Arayüz (Tarih Seçim ve Sorgulama Ekranı)
def create_main_screen():
    def fetch_and_display():
        start_date = start_cal.get_date()
        end_date = end_cal.get_date()
        if start_date > end_date:
            messagebox.showerror("Hata", "Başlangıç tarihi bitiş tarihinden sonra olamaz.")
            return

        start_date = datetime.combine(start_date, datetime.min.time())
        end_date = datetime.combine(end_date, datetime.max.time())
        events = get_events(start_date, end_date)
        plot_event_levels(events)

    def save_events_to_word():
        start_date = start_cal.get_date()
        end_date = end_cal.get_date()
        if start_date > end_date:
            messagebox.showerror("Hata", "Başlangıç tarihi bitiş tarihinden sonra olamaz.")
            return

        start_date = datetime.combine(start_date, datetime.min.time())
        end_date = datetime.combine(end_date, datetime.max.time())
        events = get_events(start_date, end_date)
        save_to_word(events)

    main_screen = tk.Toplevel(root)
    main_screen.configure(bg="black")
    main_screen.title("Olay Görüntüleyici")
    main_screen.geometry("400x300")

    tk.Label(main_screen, text="Başlangıç Tarihi", bg="black", fg="pink").pack(pady=5)
    start_cal = DateEntry(main_screen, width=12, background='darkblue', foreground='white', borderwidth=2)
    start_cal.pack(pady=5)

    tk.Label(main_screen, text="Bitiş Tarihi", bg="black", fg="pink").pack(pady=5)
    end_cal = DateEntry(main_screen, width=12, background='darkblue', foreground='white', borderwidth=2)
    end_cal.pack(pady=5)

    button1 = tk.Button(main_screen, text="Grafik Oluştur", command=fetch_and_display, bg="pink", fg="black")
    button1.pack(pady=10)

    button2 = tk.Button(main_screen, text="Word'e Kaydet", command=save_events_to_word, bg="pink", fg="black")
    button2.pack(pady=10)

    # Hover Renk Değişimi
    for button in [button1, button2]:
        button.bind("<Enter>", lambda e: e.widget.config(bg="gray"))
        button.bind("<Leave>", lambda e: e.widget.config(bg="pink"))

# Ana Ekran
root = tk.Tk()
root.configure(bg="black")
root.title("Elektronik Delil Dersi Ödev")
root.geometry("400x300")

tk.Label(root, text="Elektronik Delil Dersi Ödev", font=("Helvetica", 16), bg="black", fg="pink").pack(pady=50)

button = tk.Button(root, text="Tarih Sorgula", command=create_main_screen, bg="pink", fg="black")
button.pack(pady=10)

# Hover Renk Değişimi
button.bind("<Enter>", lambda e: e.widget.config(bg="gray"))
button.bind("<Leave>", lambda e: e.widget.config(bg="pink"))

root.mainloop()

